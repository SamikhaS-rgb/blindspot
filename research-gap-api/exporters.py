import json
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, RGBColor
from models import SessionLocal, Finding, Report


SECTION_COLORS = {
    "gap":           colors.HexColor("#185FA5"),
    "contradiction": colors.HexColor("#993C1D"),
    "methodology":   colors.HexColor("#854F0B"),
    "suggestion":    colors.HexColor("#0F6E56"),
}


def build_pdf(job_id: str, output_path: str):
    db = SessionLocal()
    report = db.query(Report).filter_by(job_id=job_id).first()
    findings = db.query(Finding).filter_by(job_id=job_id).all()
    stats = json.loads(report.stats) if report else {}
    db.close()

    doc = SimpleDocTemplate(output_path, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []

    # Title
    story.append(Paragraph("Research Gap Analysis Report", styles["Title"]))
    story.append(Spacer(1, 0.4*cm))

    # Stats table
    stat_data = [
        ["Papers analyzed", "Gaps", "Contradictions", "Suggestions"],
        [str(stats.get("papers_analyzed","—")),
         str(stats.get("gaps_found","—")),
         str(stats.get("contradictions_found","—")),
         str(stats.get("suggestions","—"))],
    ]
    t = Table(stat_data, colWidths=[4*cm]*4)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#F1EFE8")),
        ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE",   (0,0), (-1,-1), 10),
        ("ALIGN",      (0,0), (-1,-1), "CENTER"),
        ("GRID",       (0,0), (-1,-1), 0.5, colors.HexColor("#B4B2A9")),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white]),
    ]))
    story.append(t)
    story.append(Spacer(1, 0.5*cm))

    # Summary
    story.append(Paragraph("Summary", styles["Heading2"]))
    story.append(Paragraph(report.summary if report else "", styles["Normal"]))
    story.append(Spacer(1, 0.5*cm))

    # Findings by section
    for kind, label in [("gap","Research Gaps"), ("contradiction","Contradictions"),
                         ("methodology","Methodological Weaknesses"), ("suggestion","Novel Directions")]:
        section_findings = [f for f in findings if f.kind == kind]
        if not section_findings:
            continue
        story.append(Paragraph(label, styles["Heading2"]))
        for f in section_findings:
            detail = json.loads(f.detail)
            story.append(Paragraph(f"<b>{f.title}</b>", styles["Normal"]))
            story.append(Paragraph(f.description, styles["Normal"]))
            if kind == "gap" and detail.get("populations"):
                story.append(Paragraph(f"<i>Understudied: {detail['populations']}</i>", styles["Normal"]))
            if kind == "contradiction" and detail.get("implication"):
                story.append(Paragraph(f"<i>Implication: {detail['implication']}</i>", styles["Normal"]))
            if kind == "methodology" and detail.get("better_approach"):
                story.append(Paragraph(f"<i>Better approach: {detail['better_approach']}</i>", styles["Normal"]))
            story.append(Spacer(1, 0.25*cm))

    doc.build(story)


def build_docx(job_id: str, output_path: str):
    db = SessionLocal()
    report = db.query(Report).filter_by(job_id=job_id).first()
    findings = db.query(Finding).filter_by(job_id=job_id).all()
    stats = json.loads(report.stats) if report else {}
    db.close()

    doc = Document()
    doc.add_heading("Research Gap Analysis Report", 0)

    # Stats
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["Papers analyzed","Gaps","Contradictions","Suggestions"]):
        table.cell(0,i).text = h
        table.cell(0,i).paragraphs[0].runs[0].bold = True
    for i, k in enumerate(["papers_analyzed","gaps_found","contradictions_found","suggestions"]):
        table.cell(1,i).text = str(stats.get(k,"—"))
    doc.add_paragraph()

    if report:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(report.summary)

    for kind, label in [("gap","Research Gaps"), ("contradiction","Contradictions"),
                         ("methodology","Methodological Weaknesses"), ("suggestion","Novel Directions")]:
        section_findings = [f for f in findings if f.kind == kind]
        if not section_findings:
            continue
        doc.add_heading(label, level=1)
        for f in section_findings:
            detail = json.loads(f.detail)
            p = doc.add_paragraph()
            run = p.add_run(f.title)
            run.bold = True
            doc.add_paragraph(f.description)
            if kind == "gap" and detail.get("populations"):
                doc.add_paragraph(f"Understudied: {detail['populations']}").italic = True
            if kind == "methodology" and detail.get("better_approach"):
                doc.add_paragraph(f"Better approach: {detail['better_approach']}").italic = True

    doc.save(output_path)